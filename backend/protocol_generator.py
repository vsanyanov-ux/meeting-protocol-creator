import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime
import re

def add_page_number(run):
    """Helper to add page numbers via OOXML."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def generate_docx(content: str) -> str:
    """Generate a high-end enterprise DOCX protocol with professional corporate styling."""
    doc = Document()
    
    # 0. Page Margins (ГОСТ Style)
    section = doc.sections[0]
    section.left_margin = Cm(3.0)   # Wider left for binding
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    
    # Footer - Confidential & Page Numbering
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Confidential label
    conf_run = footer_p.add_run("КОНФИДЕНЦИАЛЬНО | ")
    conf_run.font.size = Pt(9)
    conf_run.font.color.rgb = RGBColor(150, 0, 0) # Dark Red for alert
    conf_run.bold = True
    
    # Page numbering
    page_run = footer_p.add_run("Страница ")
    page_run.font.size = Pt(9)
    add_page_number(footer_p.add_run())
    
    # Global Style Adjustments
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.15
    paragraph_format.space_after = Pt(6)

    # 1. Main Title
    doc.add_paragraph() # Spacer
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run('ПРОТОКОЛ СОВЕЩАНИЯ')
    run_title.bold = True
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = RGBColor(0, 51, 102) # Corporate Navy
    
    # 3. Meta Info Header
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_meta = p_meta.add_run(f"Дата формирования: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}")
    run_meta.italic = True
    run_meta.font.size = Pt(10)
    run_meta.font.color.rgb = RGBColor(100, 100, 100)

    # Horizontal Bar Separator
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = p_sep.add_run("________________________________________________________________________________")
    run_sep.font.color.rgb = RGBColor(200, 200, 200)

    # Split content by lines
    lines = content.split('\n')
    table_data = []
    
    def finalize_table(doc_obj, data):
        if not data: return
        table = doc_obj.add_table(rows=len(data), cols=4)
        table.style = 'Table Grid'
        table.autofit = False
        
        widths = [Cm(1.0), Cm(9.0), Cm(3.2), Cm(3.3)]
        for i, width in enumerate(widths):
            table.columns[i].width = width
        
        for r_idx, row_data in enumerate(data):
            for c_idx, cell_data in enumerate(row_data):
                if c_idx < 4:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = str(cell_data)
                    # Styling for header row
                    if r_idx == 0:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:fill'), 'E9E9E9') # Light Gray background
                        tcPr.append(shd)
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in p.runs:
                                run.bold = True
                                run.font.size = Pt(10)
                    else:
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in p.runs:
                                run.font.size = Pt(10)
        doc_obj.add_paragraph("")

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # Table detection
        if line_stripped.startswith('|') and line_stripped.endswith('|'):
            if re.match(r"^[|\-\s:]+$", line_stripped):
                continue
            
            cols = [col.strip() for col in line_stripped.split('|')[1:-1]]
            if len(cols) > 4: cols = cols[:4]
            while len(cols) < 4: cols.append("-")
            table_data.append(cols)
            continue
        
        if table_data:
            finalize_table(doc, table_data)
            table_data = []

        # Heading detection
        if line_stripped.startswith('##'):
            doc.add_paragraph()
            h_text = line_stripped.replace('#', '').strip().upper()
            h_p = doc.add_paragraph()
            run_h = h_p.add_run(h_text)
            run_h.font.size = Pt(14)
            run_h.bold = True
            run_h.font.color.rgb = RGBColor(0, 51, 102)
            continue
        
        # Label detection
        match_label = re.match(r"^([\w\sа-яА-ЯёЁ]+[:–])(.*)$", line_stripped)
        if match_label:
            p = doc.add_paragraph()
            run_lbl = p.add_run(match_label.group(1))
            run_lbl.bold = True
            run_lbl.font.size = Pt(12)
            p.add_run(match_label.group(2))
        else:
            p_text = re.sub(r"\*\*(.*?)\*\*", r"\1", line_stripped)
            p_text = p_text.replace("__", "").strip()
            if p_text:
                p = doc.add_paragraph(p_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if table_data:
        finalize_table(doc, table_data)
            
    # Save directory
    temp_dir = "temp_protocols"
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)
        
    filename = f"Protocol_Enterprise_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(temp_dir, filename)
    doc.save(filepath)
    
    return filepath
